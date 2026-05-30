from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Cm


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Hittade inte stycke som börjar med: {prefix}")


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def insert_before(target, text="", style="Normal"):
    element = OxmlElement("w:p")
    target._p.addprevious(element)
    paragraph = Paragraph(element, target._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_table_before(document, target, rows, widths):
    table = document.add_table(rows=0, cols=len(widths))
    table.style = "Table Grid"
    table.autofit = False

    for row_index, values in enumerate(rows):
        row = table.add_row()
        if row_index == 0:
            header_props = row._tr.get_or_add_trPr()
            header_props.append(OxmlElement("w:tblHeader"))

        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.width = Cm(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = "Normal"
            paragraph.paragraph_format.space_after = Cm(0)
            run = paragraph.add_run(value)
            if row_index == 0:
                run.bold = True
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9E2F3")
                cell._tc.get_or_add_tcPr().append(shading)

    target._p.addprevious(table._tbl)
    return table


def add_figure_before(target, image_path, caption):
    paragraph = insert_before(target)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(15.7))
    insert_before(target, caption, "Bildtext")


def main():
    if len(sys.argv) != 4:
        raise SystemExit(
            "Användning: revise-report-route-filter.py INPUT.docx OUTPUT.docx IMAGE.png"
        )

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    image_path = Path(sys.argv[3])

    document = Document(input_path)

    set_text(
        find_paragraph(document, "Detta examensarbete handlar om utveckling"),
        "Detta examensarbete handlar om utveckling och utvärdering av en mobil "
        "samåkningsprototyp. Arbetet fokuserar på positionsbaserad matchning av "
        "resor, Socket.IO-baserade uppdateringar, kontroll vid parallella "
        "godkännanden av intresseanmälningar och användarnas uppfattning av "
        "tillitsrelaterade designval. En prototyp kallad CatchMe har tagits fram. "
        "I prototypen kan användaren välja en tänkt startpunkt och målpunkt samt "
        "maximalt tillåtet avstånd till vardera punkten. Koordinater används för "
        "att filtrera och sortera matchande resor direkt i databasen.",
    )
    set_text(
        find_paragraph(document, "Prototypen utvärderades genom mätningar"),
        "Prototypen utvärderades genom mätningar av svarstid vid ökande "
        "datamängder, ett kompletterande funktionstest av ruttfiltret, en "
        "jämförelse mellan filtrering i SQL och filtrering i applikationslagret, "
        "tester av parallella godkännandeanrop, tester av Socket.IO-baserade "
        "uppdateringar i lokal testmiljö samt en enkätstudie med 29 deltagare. "
        "Funktionstestet visade att endast resor inom både vald startradie och "
        "vald målradie returnerades. I det största jämförelsefallet, med "
        "1 000 003 mätta rader, var medeltiden 584,55 ms för SQL-varianten och "
        "638,45 ms för applikationsvarianten.",
    )
    set_text(
        find_paragraph(document, "Slutsatsen är att positionsbaserad"),
        "Slutsatsen är att ruttfiltrering med valda avstånd vid både start och "
        "mål, databaserad kontroll vid godkännande av intresseanmälningar och "
        "Socket.IO-baserade uppdateringar kunde kombineras i den implementerade "
        "samåkningsprototypen. Resultaten är dock begränsade till en lokal "
        "testmiljö och en mindre användarstudie. Jämförelsetestet visar därför "
        "inte generell produktionsprestanda, utan resultat under den använda "
        "lokala mätuppställningen.",
    )
    set_text(
        find_paragraph(document, "Nyckelord:"),
        "Nyckelord: samåkning, ruttfilter, positionsbaserad matchning, "
        "Socket.IO, samtidighetskontroll, godkännandeflöde.",
    )

    set_text(
        find_paragraph(document, "This bachelor thesis concerns"),
        "This bachelor thesis concerns the development and evaluation of a "
        "mobile ride-sharing prototype. The prototype allows a user to choose "
        "an intended origin and destination and specify a maximum distance for "
        "each point. Coordinates are used to filter and rank matching trips "
        "directly in the database. The work also addresses Socket.IO-based "
        "updates, concurrent approvals of interest requests, and users' "
        "perception of trust-related design choices.",
    )
    set_text(
        find_paragraph(document, "The prototype was evaluated"),
        "The prototype was evaluated through response-time measurements, a "
        "functional test of the route filter, a comparison between SQL-side and "
        "application-side filtering, tests of concurrent approval attempts, "
        "Socket.IO tests in a local environment, and a questionnaire study with "
        "29 participants. The functional test confirmed that only trips within "
        "both selected distance limits were returned. For the largest comparison "
        "case, with 1,000,003 measured rows, the mean time was 584.55 ms for "
        "SQL-side filtering and 638.45 ms for application-side filtering.",
    )
    set_text(find_paragraph(document, "in the tested scenario."), "")
    set_text(
        find_paragraph(document, "The conclusion is that position-based"),
        "The conclusion is that route filtering using selected distance limits "
        "at both origin and destination, database-based control of interest-"
        "request approvals, and Socket.IO-based updates could be combined in "
        "the implemented prototype. The reported results are limited to the "
        "local test environment and should not be generalized to production "
        "deployments.",
    )
    set_text(find_paragraph(document, ", and Socket.IO-based updates"), "")
    set_text(
        find_paragraph(document, "Keywords:"),
        "Keywords: ride-sharing, route filtering, position-based matching, "
        "Socket.IO, concurrency control, interest request approval flow.",
    )

    limitation = find_paragraph(
        document, "Testet av positionsbaserad sortering omfattade"
    )
    set_text(
        limitation,
        "Det ursprungliga testet av positionsbaserad sortering omfattade den "
        "lösning som implementerades för det vanliga flödet. Som komplettering "
        "genomfördes ett funktionstest av ruttfiltret och en jämförelse mellan "
        "SQL-filtrering och motsvarande beräkning i applikationslagret. "
        "Jämförelsen omfattar inte externa kart- eller ruttjänster och ska "
        "tolkas som en lokal jämförelse av var beräkning och filtrering utförs.",
    )

    ui_paragraph = find_paragraph(document, "På startsidan visas resor")
    set_text(
        ui_paragraph,
        "På startsidan visas resor för användaren. Utöver det vanliga "
        "platsbaserade flödet kan användaren öppna ett ruttfilter, välja "
        "önskad startpunkt och målpunkt och ange maximal distans separat för "
        "upphämtning respektive avlämning. Standardvärdet 0,5 km motsvarar "
        "500 meter. För matchade resor visar resekortet avståndet vid både "
        "start och mål.",
    )
    finished_paragraph = find_paragraph(
        document, "Den färdiga prototypen resulterade"
    )
    set_text(
        finished_paragraph,
        "Den färdiga prototypen resulterade i en mobil samåkningsapplikation "
        "där användare kan skapa resor, visa resor och bildposter i ett "
        "gemensamt flöde, filtrera resor mot en tänkt rutt med två "
        "avståndsgränser, skicka intresseanmälningar, hantera inkomna "
        "intresseanmälningar och kommunicera via meddelanden efter accepterad "
        "intresseanmälan.",
    )

    implementation_target = find_paragraph(
        document, "Hantering av godkännande och dataintegritet"
    )
    insert_before(
        implementation_target,
        "Ruttfilter med valbart avstånd vid start och mål",
        "Rubrik 2 numrerad",
    )
    insert_before(
        implementation_target,
        "Som komplettering till det ordinarie feedflödet implementerades ett "
        "ruttfilter. Användaren väljer en tänkt startpunkt och en tänkt "
        "målpunkt via platsförslag och anger två maximala avstånd: ett till "
        "upphämtningspunkten och ett från avlämningspunkten. Förvalet är "
        "0,5 km för vardera punkten.",
    )
    insert_before(
        implementation_target,
        "Frontend skickar de valda koordinaterna och radierna till endpointen "
        "GET /post/route-search. Backend beräknar pickup_distance mellan den "
        "valda startpunkten och varje resas startkoordinater samt "
        "destination_distance mellan den valda målpunkten och varje resas "
        "målkoordinater.",
    )
    insert_before(
        implementation_target,
        "Filtreringen utförs i SQL. En resa returneras endast när både "
        "pickup_distance är mindre än eller lika med vald startradie och "
        "destination_distance är mindre än eller lika med vald målradie. "
        "Träffarna sorteras efter summan av de två avstånden. Resans id används "
        "som sekundär sorteringsnyckel så att paginering får en stabil ordning "
        "när två resor har samma beräknade avstånd.",
    )

    measurement_target = find_paragraph(
        document, "För test av godkännande av intresseanmälningar"
    )
    insert_before(
        measurement_target,
        "Kompletterande test av ruttfilter och jämförelse",
        "Rubrik 3 numrerad",
    )
    insert_before(
        measurement_target,
        "Ruttfiltrets funktionskrav testades med fyra tillfälligt skapade "
        "resor och radien 0,5 km vid både start och mål. Två resor placerades "
        "inom båda radierna. En resa placerades utanför startradien men inom "
        "målradien, och en resa placerades inom startradien men utanför "
        "målradien. Testet godkändes endast om de två första resorna "
        "returnerades i sorterad ordning och båda övriga exkluderades.",
    )
    insert_before(
        measurement_target,
        "För att lyfta utvärderingen jämfördes två lösningar med samma "
        "koordinater, två radier och sorteringsregel. I SQL-varianten gjordes "
        "avståndsberäkning, filtrering och sortering i databasen, motsvarande "
        "den implementerade endpointens princip. I applikationsvarianten "
        "hämtades koordinatrader från databasen och samma beräkning, filtrering "
        "och sortering utfördes i Node.js. Fyra datamängder mättes tio gånger "
        "vardera. Jämförelsen omfattar inte externa kartanrop och isolerar "
        "placeringen av beräkningen i den lokala testmiljön.",
    )

    results_target = find_paragraph(
        document, "Resultat från test av godkännande av intresseanmälningar"
    )
    insert_before(
        results_target,
        "Resultat från kompletterande test av ruttfilter",
        "Rubrik 3 numrerad",
    )
    insert_before(
        results_target,
        "Funktionstestet kördes med 0,5 km som högsta tillåtna avstånd vid "
        "både vald startpunkt och vald målpunkt. Testet godkändes: de två "
        "resorna som låg inom båda radierna returnerades, medan resan utanför "
        "startradien och resan utanför målradien exkluderades.",
    )
    insert_table_before(
        document,
        results_target,
        [
            ["Testresa", "Startavstånd (km)", "Målavstånd (km)", "Utfall"],
            ["MATCH_NEAREST", "0,080", "0,090", "Returnerades"],
            ["MATCH_SECOND", "0,230", "0,310", "Returnerades"],
            ["OUTSIDE_PICKUP", "0,650", "0,100", "Exkluderades"],
            ["OUTSIDE_DESTINATION", "0,100", "0,650", "Exkluderades"],
        ],
        [5.0, 3.6, 3.6, 3.6],
    )
    insert_before(
        results_target,
        "Tabell 6.1a: Funktionstest av ruttfilter med avståndsgräns vid både "
        "start och mål.",
        "Bildtext",
    )
    add_figure_before(
        results_target,
        image_path,
        "Figur 6.17a: Filterdialog för vald startpunkt, målpunkt och separata "
        "avståndsgränser.",
    )
    insert_before(
        results_target,
        "Jämförelse mellan SQL-filtrering och applikationsfiltrering",
        "Rubrik 3 numrerad",
    )
    insert_before(
        results_target,
        "Tabell 6.1b redovisar den kompletterande jämförelsen. Båda "
        "varianterna returnerade samma sorterade urval och använde högst "
        "0,5 km vid vardera punkt. Tabellen visar direkt mätt tid för de två "
        "beräkningsplaceringarna i den lokala databasmiljön.",
    )
    insert_table_before(
        document,
        results_target,
        [
            [
                "Mätta rader",
                "SQL min",
                "SQL max",
                "SQL medel",
                "App min",
                "App max",
                "App medel",
                "App/SQL",
            ],
            ["10 003", "5,73", "6,68", "6,11", "6,35", "7,90", "7,27", "1,19x"],
            [
                "50 003",
                "26,12",
                "168,70",
                "41,12",
                "26,16",
                "99,91",
                "34,52",
                "0,84x",
            ],
            [
                "100 003",
                "51,49",
                "57,76",
                "54,17",
                "50,71",
                "54,99",
                "53,28",
                "0,98x",
            ],
            [
                "1 000 003",
                "563,56",
                "612,46",
                "584,55",
                "583,23",
                "744,61",
                "638,45",
                "1,09x",
            ],
        ],
        [2.4, 2.0, 2.0, 2.2, 2.0, 2.0, 2.2, 1.6],
    )
    insert_before(
        results_target,
        "Tabell 6.1b: Jämförelse av svarstid i ms för SQL-filtrering och "
        "filtrering i applikationslagret, tio körningar per datamängd.",
        "Bildtext",
    )
    insert_before(
        results_target,
        "Resultaten visar inte en stabil tidsfördel för de mindre "
        "datamängderna. Vid 1 000 003 mätta rader var medeltiden lägre för "
        "SQL-filtrering (584,55 ms) än för applikationsfiltrering (638,45 ms). "
        "Eftersom mätningen genomfördes lokalt och en utstickare förekom i "
        "SQL-mätningen vid 50 003 rader ska jämförelsen tolkas försiktigt och "
        "inte generaliseras till produktionsdrift.",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
